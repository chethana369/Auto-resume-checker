import streamlit as st
import pandas as pd
import PyPDF2
import docx

# Set page config as the very first Streamlit command
st.set_page_config(page_title="Resume Relevance Checker", layout="wide")

# Stylish pastel UI
st.markdown("""
<style>
body, [data-testid="stAppViewContainer"], .stApp {
    background-color: #e5eaf5 !important;  /* Main background */
    color: #232323 !important;             /* Dark text */
}
[data-testid="stSidebar"] {
    background: linear-gradient(160deg, #a0d2eb 0%, #d0bdf4 90%);
}
h1, h2, h3 {
    color: #2e4057 !important; /* Dark blue for headers */
    font-weight: 700;
}
.stFileUploaderDropzone {
    background-color: #a0d2eb !important;
    border: 2px solid #d0bdf4 !important;
    border-radius: 12px;
}
.stButton>button {
    background: linear-gradient(90deg, #d0bdf4 0%, #a0d2eb 100%);
    color: #232323;
    font-weight: 700;
    border-radius: 9px;
    padding: 12px 30px;
    border: none;
    box-shadow: 0 2px 8px rgba(160, 210, 235, 0.25);
    transition: background 0.3s ease;
}
.stButton>button:hover {
    background: linear-gradient(90deg, #a0d2eb 0%, #d0bdf4 100%);
}
textarea, .stTextArea textarea {
    background-color: #d0bdf4 !important;
    color: #232323 !important;
    font-size: 15px;
    border-radius: 9px;
}
.stDataFrame {
    background-color: #e5eaf5 !important;
    color: #232323 !important;
    border-radius: 10px;
    font-weight: 600;
}
.stAlert, [data-testid="stNotificationContent"] {
    background-color: #d0bdf4 !important;
    color: #232323 !important;
    border-radius: 10px;
}
</style>
""", unsafe_allow_html=True)

st.title("🎯 Resume Relevance Check System")
st.write("Upload job description and resumes to get relevance scores")

# Sidebar for job description
with st.sidebar:
    st.header("Job Description")
    uploaded_file = st.file_uploader("Upload Job Description file", type=["txt", "pdf", "docx"])
    job_description = ""
    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            job_description = text
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            job_description = text
        else:
            job_description = uploaded_file.read().decode("utf-8")
        st.text_area("Job Description Preview", job_description, height=300)
        if job_description.strip() != "":
            st.session_state['jd_processed'] = job_description
    else:
        job_description = st.text_area("Paste Job Description", height=300)
        if job_description.strip() != "":
            st.session_state['jd_processed'] = job_description

# Helper functions for resume parsing and scoring
def extract_text_from_upload(uploaded_file):
    if uploaded_file.type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return uploaded_file.read().decode("utf-8")

def calculate_match_score(resume_text, jd_text):
    jd_keywords = set(jd_text.lower().split())
    resume_keywords = set(resume_text.lower().split())
    matched = jd_keywords & resume_keywords
    score = int(len(matched) / max(len(jd_keywords), 1) * 100)
    verdict = "Relevant" if score > 50 else "Not Relevant"
    missing_skills = list(jd_keywords - resume_keywords)
    return score, verdict, missing_skills

# Main area for resume upload and analysis
if 'jd_processed' in st.session_state and st.session_state['jd_processed'].strip() != "":
    st.header("Upload Resumes")
    uploaded_files = st.file_uploader("Choose resume files", type=['pdf', 'docx'], accept_multiple_files=True)
    
    if uploaded_files and st.button("Analyze Resumes"):
        results = []
        progress_bar = st.progress(0)

        for i, file in enumerate(uploaded_files):
            resume_text = extract_text_from_upload(file)
            score, verdict, missing_skills = calculate_match_score(
                resume_text, st.session_state['jd_processed']
            )
            results.append({
                'Resume': file.name,
                'Score': f"{score}/100",
                'Verdict': verdict,
                'Missing Skills': ', '.join(missing_skills[:5])
            })
            progress_bar.progress((i + 1) / len(uploaded_files))

        df = pd.DataFrame(results)
        st.dataframe(df, use_container_width=True)

        csv = df.to_csv(index=False)
        st.download_button("Download Results", data=csv, file_name="results.csv", mime="text/csv")
else:
    st.info("👉 Please upload or paste a job description first")
